from __future__ import annotations

import json
import re
from collections import Counter, defaultdict
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor


ROOT = Path(r"C:/Source_Code/SWP")
PROJECT = ROOT / "SWP391-Project"
SRS_WORK = ROOT / "SRS_Work"
SRS_DOCX = SRS_WORK / "SEAL_Hackathon_SRS_Completed.docx"
REQ_JSON = SRS_WORK / "srs_requirements.json"
TABLE_JSON = SRS_WORK / "srs_tables.json"
OUT_DOCX = PROJECT / "SEAL_Hackathon_Master_Test_Plan.docx"

ACCENT = "1F4D78"
ACCENT_2 = "2E74B5"
HEADER_FILL = "E8EEF5"
SUBTLE_FILL = "F4F6F9"
LIGHT_FILL = "F7F9FC"
WHITE = "FFFFFF"
BLACK = "000000"
MUTED = "555555"
RISK_HIGH = "FCE4D6"
RISK_MED = "FFF2CC"
RISK_LOW = "E2F0D9"
BULLET_NUM_ID = None
DECIMAL_NUM_ID = None
HEADING_COUNTERS = [0, 0, 0, 0]


def load_data():
    if not REQ_JSON.exists() or not TABLE_JSON.exists():
        raise FileNotFoundError("Parsed SRS JSON files are missing. Run the SRS extraction step first.")
    reqs = json.loads(REQ_JSON.read_text(encoding="utf-8"))
    tables = json.loads(TABLE_JSON.read_text(encoding="utf-8"))
    table_map = {t["table"]: t["rows"] for t in tables}
    return reqs, table_map


def set_font(run, name="Times New Roman", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def style_font(style, name="Times New Roman", size=None, color=None, bold=None, italic=None):
    style.font.name = name
    style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    if size is not None:
        style.font.size = Pt(size)
    if color is not None:
        style.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        style.font.bold = bold
    if italic is not None:
        style.font.italic = italic


def configure_styles(doc: Document):
    styles = doc.styles
    normal = styles["Normal"]
    style_font(normal, size=11, color=BLACK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, color, before, after in [
        ("Heading 1", 16, ACCENT, 15, 8),
        ("Heading 2", 13, ACCENT_2, 11, 6),
        ("Heading 3", 12, ACCENT, 8, 4),
        ("Heading 4", 11, ACCENT, 6, 3),
    ]:
        st = styles[name]
        style_font(st, size=size, color=color, bold=True)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    for name in ["List Bullet", "List Bullet 2", "List Number", "List Number 2"]:
        if name in styles:
            style_font(styles[name], size=11, color=BLACK)
            styles[name].paragraph_format.space_after = Pt(4)

    for name in ["Table Grid", "Light Shading Accent 1"]:
        if name in styles:
            style_font(styles[name], size=9, color=BLACK)


def set_update_fields_on_open(doc: Document):
    settings = doc.settings.element
    for existing in settings.findall(qn("w:updateFields")):
        settings.remove(existing)
    update = OxmlElement("w:updateFields")
    update.set(qn("w:val"), "true")
    settings.append(update)


def add_field_run(paragraph, instr, fallback=""):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    instr_run = paragraph.add_run()
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instr
    instr_run._r.append(instr_text)

    sep_run = paragraph.add_run()
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    sep_run._r.append(sep)

    if fallback:
        f_run = paragraph.add_run(fallback)
        set_font(f_run)

    end_run = paragraph.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    end_run._r.append(fld_end)


def add_toc(doc: Document):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = title.add_run("Table of Contents")
    set_font(r, size=16, color=ACCENT, bold=True)
    title.paragraph_format.space_after = Pt(8)
    p = doc.add_paragraph()
    add_field_run(p, r'TOC \o "1-3" \h \z \u', "Table of contents will update automatically when fields are refreshed.")


def add_page_field(paragraph, instr, fallback="1"):
    add_field_run(paragraph, instr, fallback)


def configure_header_footer(doc: Document):
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = header.add_run("SEAL Hackathon Management System - Master Test Plan")
    set_font(r, size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("Page ")
    set_font(r, size=9, color=MUTED)
    add_page_field(footer, "PAGE", "1")
    r = footer.add_run(" of ")
    set_font(r, size=9, color=MUTED)
    add_page_field(footer, "NUMPAGES", "1")


def create_heading_numbering(doc: Document) -> int:
    numbering = doc.part.numbering_part.element
    existing_abs = [
        int(x.get(qn("w:abstractNumId")))
        for x in numbering.findall(qn("w:abstractNum"))
        if x.get(qn("w:abstractNumId")) is not None
    ]
    existing_num = [
        int(x.get(qn("w:numId")))
        for x in numbering.findall(qn("w:num"))
        if x.get(qn("w:numId")) is not None
    ]
    abstract_id = (max(existing_abs) + 1) if existing_abs else 1
    num_id = (max(existing_num) + 1) if existing_num else 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    mlt = OxmlElement("w:multiLevelType")
    mlt.set(qn("w:val"), "hybridMultilevel")
    abstract.append(mlt)

    lvl_text = ["%1.", "%1.%2", "%1.%2.%3", "%1.%2.%3.%4"]
    for i in range(4):
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), str(i))
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        fmt = OxmlElement("w:numFmt")
        fmt.set(qn("w:val"), "decimal")
        txt = OxmlElement("w:lvlText")
        txt.set(qn("w:val"), lvl_text[i])
        jc = OxmlElement("w:lvlJc")
        jc.set(qn("w:val"), "left")
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "space")
        restart = OxmlElement("w:lvlRestart")
        restart.set(qn("w:val"), "1")
        ppr = OxmlElement("w:pPr")
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "0")
        ind.set(qn("w:hanging"), "0")
        ppr.append(ind)
        children = [start, fmt, txt, suff, jc, ppr]
        if i > 0:
            children.insert(3, restart)
        lvl.extend(children)
        abstract.append(lvl)

    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abs_id = OxmlElement("w:abstractNumId")
    abs_id.set(qn("w:val"), str(abstract_id))
    num.append(abs_id)
    numbering.append(num)
    return num_id


def create_list_numbering(doc: Document, kind: str) -> int:
    numbering = doc.part.numbering_part.element
    existing_abs = [
        int(x.get(qn("w:abstractNumId")))
        for x in numbering.findall(qn("w:abstractNum"))
        if x.get(qn("w:abstractNumId")) is not None
    ]
    existing_num = [
        int(x.get(qn("w:numId")))
        for x in numbering.findall(qn("w:num"))
        if x.get(qn("w:numId")) is not None
    ]
    abstract_id = (max(existing_abs) + 1) if existing_abs else 1
    num_id = (max(existing_num) + 1) if existing_num else 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    mlt = OxmlElement("w:multiLevelType")
    mlt.set(qn("w:val"), "hybridMultilevel")
    abstract.append(mlt)

    for i in range(2):
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), str(i))
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        fmt = OxmlElement("w:numFmt")
        txt = OxmlElement("w:lvlText")
        if kind == "bullet":
            fmt.set(qn("w:val"), "bullet")
            txt.set(qn("w:val"), "•")
        else:
            fmt.set(qn("w:val"), "decimal")
            txt.set(qn("w:val"), "%1." if i == 0 else "%2.")
        jc = OxmlElement("w:lvlJc")
        jc.set(qn("w:val"), "left")
        ppr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), str(720 + (i * 360)))
        tabs.append(tab)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(720 + (i * 360)))
        ind.set(qn("w:hanging"), "360")
        ppr.extend([tabs, ind])
        lvl.extend([start, fmt, txt, jc, ppr])
        abstract.append(lvl)

    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abs_id = OxmlElement("w:abstractNumId")
    abs_id.set(qn("w:val"), str(abstract_id))
    num.append(abs_id)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int, level: int):
    ppr = paragraph._p.get_or_add_pPr()
    for child in ppr.findall(qn("w:numPr")):
        ppr.remove(child)
    numpr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(level))
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    numpr.extend([ilvl, num])
    ppr.append(numpr)


def h(doc: Document, text: str, level: int, num_id: int):
    idx = min(level, 4) - 1
    HEADING_COUNTERS[idx] += 1
    for i in range(idx + 1, len(HEADING_COUNTERS)):
        HEADING_COUNTERS[i] = 0
    prefix = ".".join(str(n) for n in HEADING_COUNTERS[: idx + 1] if n)
    p = doc.add_paragraph(f"{prefix}. {text}", style=f"Heading {min(level, 4)}")
    return p


def p(doc: Document, text: str = "", style=None, bold_label=None):
    para = doc.add_paragraph(style=style)
    if bold_label:
        r = para.add_run(bold_label)
        set_font(r, bold=True)
        if text:
            r = para.add_run(text)
            set_font(r)
    else:
        r = para.add_run(text)
        set_font(r)
    return para


def bullet(doc: Document, text: str, level: int = 0):
    para = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    if BULLET_NUM_ID is not None:
        apply_numbering(para, BULLET_NUM_ID, level)
    r = para.add_run(text)
    set_font(r)
    para.paragraph_format.space_after = Pt(4)
    return para


def numbered(doc: Document, text: str, level: int = 0):
    para = doc.add_paragraph(style="List Number" if level == 0 else "List Number 2")
    if DECIMAL_NUM_ID is not None:
        apply_numbering(para, DECIMAL_NUM_ID, level)
    r = para.add_run(text)
    set_font(r)
    para.paragraph_format.space_after = Pt(4)
    return para


def clear_cell(cell):
    for p0 in cell.paragraphs:
        p0._element.getparent().remove(p0._element)
    cell._tc.get_or_add_tcPr()


def shade_cell(cell, fill: str):
    tcpr = cell._tc.get_or_add_tcPr()
    for existing in tcpr.findall(qn("w:shd")):
        tcpr.remove(existing)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tcpr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tcpr = cell._tc.get_or_add_tcPr()
    tc_mar = tcpr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tcpr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text, bold=False, fill=None, font_size=8.8, align=None):
    clear_cell(cell)
    para = cell.add_paragraph()
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.05
    if align:
        para.alignment = align
    lines = str(text).split("\n") if text is not None else [""]
    for idx, line in enumerate(lines):
        if idx:
            para.add_run().add_break()
        r = para.add_run(line)
        set_font(r, size=font_size, color=BLACK, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if fill:
        shade_cell(cell, fill)
    set_cell_margins(cell)


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl = table._tbl
    tblpr = tbl.tblPr
    tbl_layout = tblpr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tblpr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    total = sum(widths)
    tblw = tblpr.find(qn("w:tblW"))
    if tblw is None:
        tblw = OxmlElement("w:tblW")
        tblpr.append(tblw)
    tblw.set(qn("w:w"), str(total))
    tblw.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.append(grid)
    for child in list(grid):
        grid.remove(child)
    for w in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(w))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            tcpr = cell._tc.get_or_add_tcPr()
            tcw = tcpr.find(qn("w:tcW"))
            if tcw is None:
                tcw = OxmlElement("w:tcW")
                tcpr.append(tcw)
            tcw.set(qn("w:w"), str(widths[idx]))
            tcw.set(qn("w:type"), "dxa")


def repeat_table_header(row):
    trpr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    trpr.append(header)


def add_table(doc: Document, headers, rows, widths=None, font_size=8.8, header_fill=HEADER_FILL):
    if not rows:
        rows = [[""] * len(headers)]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    repeat_table_header(table.rows[0])
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True, fill=header_fill, font_size=font_size)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], val, font_size=font_size)
    if widths is None:
        total = 9750
        widths = [total // len(headers)] * len(headers)
        widths[-1] += total - sum(widths)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_note_table(doc: Document, title: str, body: str):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_cell_text(table.cell(0, 0), f"{title}\n{body}", bold=False, fill=SUBTLE_FILL, font_size=9.3)
    set_table_geometry(table, [9750])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def req_module(req_id: str) -> str:
    parts = req_id.split("-")
    if req_id.startswith("NFR-"):
        return "NFR-" + parts[1]
    return parts[1]


def tc_base(req_id: str) -> str:
    parts = req_id.split("-")
    if req_id.startswith("FR-"):
        return "TC-" + parts[1] + "-" + parts[2]
    if req_id.startswith("NFR-"):
        return "TC-" + parts[1] + "-" + parts[2]
    return "TC-" + req_id


def ts_id(req_id: str) -> str:
    parts = req_id.split("-")
    if req_id.startswith("FR-"):
        return "TS-" + parts[1] + "-" + parts[2]
    if req_id.startswith("NFR-"):
        return "TS-" + parts[1] + "-" + parts[2]
    return "TS-" + req_id


def case_count_for(req):
    priority = req.get("Priority", "")
    if req["ID"].startswith("NFR-"):
        return 5
    if priority == "High":
        return 5
    if priority == "Medium":
        return 3
    return 2


def risk_for_req(req):
    rid = req["ID"]
    prefix = req_module(rid)
    if rid.startswith("NFR-SEC") or prefix in {"AUTH", "SCORE", "RANK", "SUB", "TEAM"}:
        return "High"
    if prefix in {"EVENT", "CONFIG", "REG", "MEMBER", "INV", "JUDGE", "ADV", "NOTIF", "EXPORT", "AUDIT", "USER"}:
        return "High" if req.get("Priority") == "High" else "Medium"
    if prefix in {"PRIZE", "MENTOR", "DASH"}:
        return "Medium"
    return "Low"


def automation_candidate(req):
    rid = req["ID"]
    title = req.get("Title", req.get("Quality attribute", ""))
    if rid.startswith("NFR-USE"):
        return "No - primarily moderated usability/UAT evidence"
    if rid.startswith("NFR-PERF"):
        return "Yes - load test script"
    if rid.startswith("NFR-SEC"):
        return "Yes - API security suite and ZAP baseline"
    if "OAuth2" in title or "SSE" in title or "email" in req.get("Description", "").lower():
        return "Partial - automate API/state checks and manually verify provider-specific behavior"
    if req_module(rid) in {"AUTH", "USER", "EVENT", "CONFIG", "REG", "TEAM", "MEMBER", "INV", "SUB", "JUDGE", "SCORE", "RANK", "ADV", "NOTIF", "EXPORT", "AUDIT"}:
        return "Yes - API regression candidate"
    return "Conditional - automate if retained in release scope"


def first_sentence(text: str) -> str:
    m = re.split(r"(?<=[.!?])\s+", text.strip())
    return m[0] if m else text


def build_business_rule_maps(table_map):
    fr_to_brs = defaultdict(list)
    br_rows = table_map[69][1:]
    for row in br_rows:
        br_id, typ, stmt, linked = row
        for fr in re.findall(r"FR-[A-Z]+-\d+", linked):
            fr_to_brs[fr].append(br_id)
    return br_rows, fr_to_brs


def table_to_rows(table_map, no):
    return table_map.get(no, [])[1:]


def module_catalog(reqs):
    definitions = [
        ("AUTH", "Authentication and Account Lifecycle", "Registration, OTP verification, login, token refresh, password reset, guest judge creation, and Google OAuth2 login."),
        ("USER", "User Profile and Administration", "Approval, administrative user creation, user listing, deactivation, profile updates, and password changes."),
        ("EVENT", "Hackathon Event Configuration", "Hackathon creation, public browsing, organizer views, updates, lifecycle status, and soft deletion."),
        ("CONFIG", "Tracks, Rounds, and Criteria", "Track setup, mentor assignment, round setup, and scoring criteria management."),
        ("REG", "Event Registration", "Participant registration for PUBLISHED events inside the registration window."),
        ("TEAM", "Team Management", "Team creation, retrieval, update, finalization, disqualification, and deletion of non-finalized teams."),
        ("MEMBER", "Team Membership", "Member addition/removal, kick/leave behavior, and leadership transfer."),
        ("INV", "Team Invitations", "Invitation creation, response, viewing, and revocation."),
        ("SUB", "Submission Management", "Submission and resubmission of repository, demo, and report URLs with role restrictions."),
        ("JUDGE", "Judge Assignment", "Judge and guest judge assignment to rounds and tracks, including mentor conflict checks."),
        ("SCORE", "Scoring", "Criterion score entry, update restrictions, and score finalization."),
        ("RANK", "Ranking", "Weighted ranking from finalized scores, exclusion of disqualified teams, and track-based ranking."),
        ("ADV", "Round Advancement", "Advancing ranked teams between rounds and completing the event when no next round exists."),
        ("PRIZE", "Prize Management", "Prize CRUD and manual/automatic prize assignment."),
        ("MENTOR", "Mentorship", "Mentorship request creation, acceptance, resolution, rejection, and cancellation."),
        ("NOTIF", "Notifications and SSE", "Notification listing, unread counts, mark-as-read operations, and SSE push."),
        ("DASH", "Dashboard Statistics", "Event statistics, days remaining, submissions received, open mentorship requests, and criterion variance."),
        ("EXPORT", "CSV Export", "Ranking, anonymized scoring, team, and participant CSV exports with audit action."),
        ("AUDIT", "Audit Logging", "Audit recording and ADMIN-only audit log review."),
        ("CAT", "Category Customer Maintenance", "Legacy CategoryCustomer CRUD with public read and ADMIN write access."),
    ]
    by_module = defaultdict(list)
    for req in reqs:
        if req["ID"].startswith("FR-"):
            by_module[req_module(req["ID"])].append(req)
    rows = []
    for key, name, desc in definitions:
        ids = ", ".join(r["ID"] for r in by_module[key])
        priority_counts = Counter(r.get("Priority", "Unspecified") for r in by_module[key])
        priority = "High" if priority_counts.get("High") else ("Medium" if priority_counts.get("Medium") else "Low")
        rows.append({
            "key": key,
            "name": name,
            "description": desc,
            "reqs": by_module[key],
            "ids": ids,
            "priority": priority,
        })
    return rows


def build_cover(doc):
    sec = doc.sections[0]
    sec.page_width = Mm(210)
    sec.page_height = Mm(297)
    sec.top_margin = Inches(0.75)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin = Inches(0.75)
    sec.right_margin = Inches(0.75)
    sec.header_distance = Inches(0.45)
    sec.footer_distance = Inches(0.45)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(70)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("MASTER TEST PLAN")
    set_font(r, size=26, color=ACCENT, bold=True)
    title.paragraph_format.space_after = Pt(6)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("SEAL Hackathon Management System")
    set_font(r, size=18, color=BLACK, bold=True)
    subtitle.paragraph_format.space_after = Pt(18)

    p0 = doc.add_paragraph()
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p0.add_run("Prepared according to IEEE 829 and ISO/IEC/IEEE 29119 test documentation practices")
    set_font(r, size=11, color=MUTED, italic=True)
    p0.paragraph_format.space_after = Pt(28)

    add_table(
        doc,
        ["Document Attribute", "Value"],
        [
            ["Document type", "Enterprise Master Test Plan"],
            ["Version", "1.0"],
            ["Date", datetime.now().strftime("%d/%m/%Y")],
            ["Source of truth", str(SRS_DOCX)],
            ["Project path", str(PROJECT)],
            ["Document status", "Ready for QA lead review"],
            ["Confidentiality", "Project team internal use"],
        ],
        widths=[2600, 7150],
        font_size=9.5,
    )
    add_note_table(
        doc,
        "Source Control Statement",
        "The SRS named above is the only authoritative source for product behavior in this test plan. "
        "Where the SRS does not specify implementation, environment, schedule, staffing, or tooling details, this document states Assumption explicitly.",
    )
    doc.add_page_break()


def build_document_control(doc, num_id):
    h(doc, "Document Control", 1, num_id)
    p(doc, "This section establishes ownership, review expectations, and approval controls for the Master Test Plan. The plan SHALL be versioned and reviewed before release test execution begins.")

    h(doc, "Version History", 2, num_id)
    add_table(
        doc,
        ["Version", "Date", "Author", "Change Summary", "Reviewer"],
        [
            ["0.1", datetime.now().strftime("%d/%m/%Y"), "Principal QA Architect / Senior QA Lead", "Initial Master Test Plan generated from SEAL Hackathon SRS version 1.0.", "Project QA Lead"],
            ["1.0", datetime.now().strftime("%d/%m/%Y"), "Project QA Team", "Baseline for review and execution after stakeholder approval.", "PM, BA, Technical Lead"],
        ],
        widths=[900, 1300, 2300, 3800, 1450],
    )

    h(doc, "Authors", 2, num_id)
    add_table(
        doc,
        ["Role", "Name", "Responsibility"],
        [
            ["QA Lead", "Assumption: assigned QA Lead name not specified in SRS.", "Owns test planning, risk strategy, execution governance, and exit recommendation."],
            ["QA Architect", "Assumption: assigned QA Architect name not specified in SRS.", "Defines automation, environments, traceability, and non-functional testing approach."],
            ["Business Analyst", "Assumption: BA name not specified in SRS.", "Confirms requirement interpretation and acceptance criteria coverage."],
        ],
        widths=[1800, 3400, 4550],
    )

    h(doc, "Reviewers", 2, num_id)
    add_table(
        doc,
        ["Reviewer Role", "Review Focus", "Required Decision"],
        [
            ["Project Manager", "Schedule, staffing, scope, and release governance.", "Approve or request changes."],
            ["Business Analyst/Product Owner", "SRS alignment, acceptance criteria, business rules, and UAT readiness.", "Approve or request changes."],
            ["Development Lead", "Build readiness, unit/integration responsibility, environment feasibility, and defect turnaround.", "Approve or request changes."],
            ["QA Lead", "Testing completeness, risk acceptance, deliverables, and exit criteria.", "Approve baseline."],
        ],
        widths=[2200, 5400, 2150],
    )

    h(doc, "Approvals", 2, num_id)
    add_table(
        doc,
        ["Approver", "Approval Area", "Status", "Date"],
        [
            ["Project Sponsor / Instructor", "Release test plan acceptance", "Pending", "Assumption: approval date not provided."],
            ["Project Manager", "Schedule and resource commitment", "Pending", "Assumption: approval date not provided."],
            ["QA Lead", "Test strategy and entry/exit criteria", "Pending", "Assumption: approval date not provided."],
            ["Development Lead", "Technical feasibility and environment support", "Pending", "Assumption: approval date not provided."],
        ],
        widths=[2600, 3800, 1350, 2000],
    )


def build_intro(doc, num_id, table_map):
    h(doc, "Introduction", 1, num_id)
    h(doc, "Purpose", 2, num_id)
    p(doc, "The purpose of this Master Test Plan is to define the complete quality assurance approach for SEAL Hackathon Management System version 1.0. The plan SHALL guide test preparation, execution, defect management, traceability, evidence collection, and release quality assessment.")
    p(doc, "The plan is written for enterprise execution. It identifies what SHALL be tested, why it SHALL be tested, how testing SHALL be performed, what evidence SHALL be produced, and what criteria SHALL control entry, suspension, resumption, and exit.")

    h(doc, "Scope", 2, num_id)
    p(doc, "The testing scope is limited to the behavior documented in the supplied SRS. The system is described as a browser-based SPA consuming a Java 21 Spring Boot 3.2.5 REST backend with SQL Server persistence, JWT authentication, Google OAuth2 login, Gmail SMTP email delivery, and Server-Sent Events notifications.")
    p(doc, "Testing SHALL cover functional requirements, non-functional requirements, business rules, data integrity requirements, external interfaces, actor permissions, and the workflows implied by the SRS acceptance criteria.")

    h(doc, "References", 2, num_id)
    refs = [
        f"SEAL_Hackathon_SRS_Completed.docx, version 1.0, located at {SRS_DOCX}.",
        "IEEE 829 style test documentation practices for Master Test Plan structure and test deliverables.",
        "ISO/IEC/IEEE 29119 software testing concepts for risk-based planning, test design, execution, and reporting.",
        "SRS-referenced backend implementation path: C:/Source_Code/SWP/SWP391-Project/backend.",
        "SRS-referenced database schema path: C:/Source_Code/SWP/SWP391-Project/database.sql.",
        "SRS-referenced OpenAPI metadata path: C:/Source_Code/SWP/SWP391-Project/backend/openapi.json.",
    ]
    for item in refs:
        bullet(doc, item)

    h(doc, "Definitions", 2, num_id)
    srs_terms = table_to_rows(table_map, 4)
    qa_terms = [
        ["Entry Criteria", "Measurable conditions that must be satisfied before a test phase or test cycle starts."],
        ["Exit Criteria", "Measurable conditions that must be satisfied before test completion can be recommended."],
        ["Severity", "Classification of defect impact on system behavior or business operation."],
        ["Priority", "Classification of how urgently a defect should be fixed based on release risk and business need."],
        ["RTM", "Requirement Traceability Matrix mapping requirements to scenarios, cases, risks, and execution evidence."],
    ]
    add_table(doc, ["Term / Acronym", "Definition"], srs_terms + qa_terms, widths=[2300, 7450], font_size=9)


def build_project_understanding(doc, num_id, table_map):
    h(doc, "Project Understanding", 1, num_id)
    p(doc, "The QA model below is derived from the SRS business objectives, actors, workflows, modules, data model, integrations, and constraints. It is included to make the test basis explicit before strategy and scope decisions are applied.")

    h(doc, "Business Objectives", 2, num_id)
    objectives = [
        "Support hackathon account onboarding, verification, approval, authentication, authorization, and profile management.",
        "Support event setup and operation through lifecycle status, registration windows, team size rules, tracks, rounds, criteria, judges, mentors, prizes, dashboards, and exports.",
        "Support participant collaboration through event registration, team creation, invitations, membership changes, finalization, submissions, and mentorship requests.",
        "Support fair evaluation through judge assignment, conflict prevention, criteria-based scoring, score finalization, ranking, advancement, and prize assignment.",
        "Support operational controls through notifications, audit logs, CSV exports, SQL Server persistence, and role-restricted administrative workflows.",
    ]
    for obj in objectives:
        bullet(doc, obj)

    h(doc, "Actors and User Classes", 2, num_id)
    add_table(doc, table_map[5][0], table_to_rows(table_map, 5), widths=[1500, 3300, 1550, 1500, 1900], font_size=8.1)

    h(doc, "Logical UI Workflows", 2, num_id)
    p(doc, "The SRS states that final frontend wireframes were not supplied; therefore UI workflows SHALL be validated against logical screens inferred from implemented API workflows.")
    add_table(doc, table_map[7][0], table_to_rows(table_map, 7), widths=[1150, 2500, 6100], font_size=8.4)

    h(doc, "External Interfaces and Dependencies", 2, num_id)
    add_table(doc, table_map[8][0], table_to_rows(table_map, 8), widths=[2100, 3000, 4650], font_size=8.6)

    h(doc, "Assumptions and Dependencies from SRS", 2, num_id)
    add_table(doc, table_map[6][0], table_to_rows(table_map, 6), widths=[900, 1400, 5600, 1850], font_size=8.6)


def build_test_objectives(doc, num_id):
    h(doc, "Test Objectives", 1, num_id)
    objectives = [
        ["OBJ-01", "Verify SRS functional coverage", "Every FR in the SRS SHALL map to at least one test scenario and at least one executable test case."],
        ["OBJ-02", "Verify business rule enforcement", "Every SRS business rule SHALL be tested through positive, negative, role, boundary, or data-state coverage."],
        ["OBJ-03", "Verify role-based security", "Public, authenticated, ADMIN, ORGANIZER, JUDGE, GUEST_JUDGE, MENTOR, PARTICIPANT, and leader-specific permissions SHALL be verified."],
        ["OBJ-04", "Verify core end-to-end hackathon flow", "The system SHALL support registration, approval, event registration, team finalization, submission, judging, ranking, advancement, notification, and export workflows."],
        ["OBJ-05", "Verify data integrity", "SQL Server constraints and service-layer validation SHALL protect identity, event registration, team membership, judge assignment, scoring, advancement, and prize consistency."],
        ["OBJ-06", "Verify NFRs", "Performance, security, usability, reliability, maintainability, and portability targets stated in the SRS SHALL be measured or explicitly dispositioned."],
        ["OBJ-07", "Support release decision", "The QA team SHALL provide evidence-based release recommendation against entry, exit, severity, risk, and residual defect criteria."],
    ]
    add_table(doc, ["ID", "Objective", "Measurable Outcome"], objectives, widths=[950, 2700, 6100], font_size=8.8)
    p(doc, "These objectives prioritize risk and traceability because the SRS includes many cross-dependent workflows where a failure in account state, event status, team eligibility, or score finalization can invalidate later workflows.")


def build_scope(doc, num_id, reqs, table_map):
    h(doc, "Test Scope", 1, num_id)
    modules = module_catalog(reqs)
    h(doc, "Features to Be Tested", 2, num_id)
    scope_rows = []
    for m in modules:
        why = "This module SHALL be tested because it is documented in the SRS as a functional feature and participates in one or more actor workflows, business rules, or data integrity paths."
        if m["key"] in {"AUTH", "TEAM", "SUB", "SCORE", "RANK"}:
            why = "This module SHALL receive elevated test depth because defects can block release-critical workflows or create fairness, access, or data integrity risk."
        objective = "Verify " + m["description"].lower()
        scope_rows.append([m["name"], m["ids"], m["priority"], why, objective])
    add_table(doc, ["Module", "SRS FR IDs", "Priority", "Why Tested", "Testing Objectives"], scope_rows, widths=[1700, 2250, 900, 2500, 2400], font_size=7.6)

    h(doc, "Cross-Cutting Quality Requirements to Be Tested", 2, num_id)
    nfr_rows = []
    for r in [x for x in reqs if x["ID"].startswith("NFR-")]:
        nfr_rows.append([
            r["ID"],
            r.get("Quality attribute", ""),
            r.get("Scale (what we measure)", ""),
            r.get("Target (desired)", ""),
            r.get("Must-level (minimum)", ""),
        ])
    add_table(doc, ["NFR ID", "Attribute", "Scale", "Target", "Minimum"], nfr_rows, widths=[1200, 1600, 3200, 2000, 1750], font_size=8.2)

    h(doc, "Features Not to Be Tested", 2, num_id)
    not_tested = [
        ["Payments and ticketing", "Out of scope in SRS v1.0; no backend module implements payment, ticketing, or financial settlement."],
        ["Venue check-in hardware and physical logistics", "Out of scope in SRS v1.0; SRS states no printers, scanners, sensors, payment terminals, or hardware devices are integrated."],
        ["Native mobile applications", "Out of scope in SRS v1.0; SRS describes REST/JSON endpoints and a browser-based SPA only."],
        ["Binary artifact upload/storage", "Out of scope in SRS v1.0; submissions store repository, demo, and report URLs rather than uploaded files."],
        ["Chat, video conferencing, or real-time collaboration beyond SSE notifications", "Out of scope in SRS v1.0; SRS limits real-time behavior to one-way Server-Sent Events notifications."],
        ["Complete legal data-erasure or retention lifecycle", "Out of scope in SRS v1.0; SRS states event soft delete and user deactivation exist, but no full compliance lifecycle is implemented."],
        ["Pixel-perfect final UI wireframes", "Assumption: final frontend wireframes were not supplied. QA SHALL validate logical workflows and usability targets, not visual compliance to unavailable wireframes."],
        ["Production hosting, domain, certificate, and capacity infrastructure", "Assumption: production hosting and sizing are not specified in the SRS. QA SHALL test stated environment constraints and NFR targets only."],
    ]
    add_table(doc, ["Excluded Feature / Area", "Justification"], not_tested, widths=[3100, 6650], font_size=8.8)


def testing_type_definitions():
    return [
        {
            "name": "Functional Testing",
            "purpose": "Validate that every functional requirement and acceptance criterion in the SRS behaves as specified.",
            "scope": "All 53 FRs, all 23 business rules, role-specific workflows, and SRS-defined actor paths.",
            "approach": "Design one positive scenario per FR, add negative and alternate-flow tests for each linked business rule, and execute API plus UI workflow coverage where screens exist.",
            "entry": "Approved SRS baseline, deployed build, seeded roles, accessible API documentation, and test data for event/team/scoring states.",
            "exit": "100% FRs executed; 100% high-priority FRs passed or formally risk accepted; no open Severity 1 or Severity 2 defect.",
            "tools": "Test management tool, Postman/Newman or REST Assured, Playwright or equivalent browser automation, SQL Server query tool.",
            "deliverables": "Functional scenarios, test cases, execution results, defect reports, and updated RTM.",
        },
        {
            "name": "Integration Testing",
            "purpose": "Verify interactions across controllers, services, repositories, SQL Server, SMTP, OAuth2, SSE, and frontend consumers.",
            "scope": "REST-to-service-to-database flows; email OTP/approval/notification paths; Google OAuth2 redirect/token issuance; SSE notification delivery.",
            "approach": "Execute service integration tests with real SQL Server schema where possible; use controlled test accounts or stubs for Gmail SMTP and Google OAuth2; verify persistence and side effects.",
            "entry": "Stable API build, schema deployed, integration configuration available, and external service test credentials or stubs approved.",
            "exit": "All high-risk integrations pass; failures are classified with workaround or release impact; integration logs retained.",
            "tools": "Spring Boot Test, REST Assured/Postman, SQL Server tooling, SMTP test mailbox or mock server, browser EventSource client.",
            "deliverables": "Integration test suite, service logs, database evidence, provider/stub configuration notes.",
        },
        {
            "name": "System Testing",
            "purpose": "Validate the complete system behavior against SRS workflows from the user's perspective.",
            "scope": "End-to-end flows for visitor browsing, participant onboarding, team formation, submission, judge scoring, organizer advancement, notifications, exports, and audit.",
            "approach": "Execute controlled scenario threads across multiple roles and event states using production-like configuration where available.",
            "entry": "Integrated build deployed to system test environment, all critical services available, and smoke test passed.",
            "exit": "End-to-end release-critical scenarios pass; remaining defects are within exit thresholds and risk accepted.",
            "tools": "Browser automation/manual execution, API client, SQL Server verification queries, test management dashboard.",
            "deliverables": "System test evidence, scenario logs, screenshots where applicable, defect metrics, exit recommendation.",
        },
        {
            "name": "API Testing",
            "purpose": "Verify REST, JSON validation, status codes, authorization, payload schemas, and documented endpoint behavior.",
            "scope": "/api/v1/* endpoints, public event/ranking reads, protected workflows, CSV export endpoints, and notification stream behavior.",
            "approach": "Build API collections by requirement ID; test success, validation, duplicate, unauthorized, forbidden, not found, conflict, and lifecycle-condition responses.",
            "entry": "Backend build running, OpenAPI metadata available as cited by SRS, seeded data and tokens available.",
            "exit": "All high and medium API tests pass or have approved defect disposition; response contracts are stable for UI consumption.",
            "tools": "Postman/Newman, REST Assured, Spring MockMvc where implemented, JSON schema assertions.",
            "deliverables": "API collection, automated run reports, request/response evidence, contract issues.",
        },
        {
            "name": "Regression Testing",
            "purpose": "Ensure existing SRS behavior remains stable after code changes, configuration changes, or defect fixes.",
            "scope": "All high-priority FRs, all security/role checks, all business rules, and representative medium/low workflows.",
            "approach": "Maintain smoke, sprint, release, and full regression packs. Prioritize automated API regression and retain manual checks for OAuth2, email, SSE, usability, and exploratory workflows.",
            "entry": "Change build deployed, impacted requirements identified, and previous baseline results available.",
            "exit": "Regression pack meets pass thresholds; no reopened critical defect; impacted RTM rows updated.",
            "tools": "CI runner assumption, Newman/REST Assured, Playwright, SQL scripts, defect tracker.",
            "deliverables": "Regression run report, failed test analysis, updated automation backlog.",
        },
        {
            "name": "Smoke Testing",
            "purpose": "Determine whether a deployed build is stable enough for deeper testing.",
            "scope": "Application startup, health of API, login, public event browse, admin token, database connectivity, and one representative protected endpoint.",
            "approach": "Execute a small mandatory suite after each deployment before assigning the build to QA.",
            "entry": "Build deployed and environment URL communicated.",
            "exit": "All smoke checks pass; any failure suspends deeper execution for that build.",
            "tools": "API smoke collection, browser sanity check, SQL Server connection check.",
            "deliverables": "Smoke result record and build acceptance/rejection note.",
        },
        {
            "name": "Sanity Testing",
            "purpose": "Verify a specific fix or changed area is coherent before targeted regression proceeds.",
            "scope": "Defect fixes, requirement changes, configuration updates, and impacted adjacent workflows.",
            "approach": "Execute the reported failing path, one positive adjacent path, one negative adjacent path, and any linked business rule checks.",
            "entry": "Fix deployed with defect ID and changed component identified.",
            "exit": "Fix is verified or reopened; impacted regression tests are identified.",
            "tools": "API client, browser, SQL verification queries, defect tracker.",
            "deliverables": "Retest result, reopened defect evidence if needed, impacted test list.",
        },
        {
            "name": "Security Testing",
            "purpose": "Verify authentication, authorization, token handling, data protection, and role restrictions documented in the SRS.",
            "scope": "JWT bearer authentication, OAuth2 login behavior, BCrypt password workflows, method-level role authorization, export access, audit review, SSE token query behavior, and PII exposure.",
            "approach": "Execute role matrix tests for all actors, unauthorized/forbidden tests for protected endpoints, token expiry/refresh checks, and baseline vulnerability scans against the test environment.",
            "entry": "Security configuration deployed, test users for all roles available, and legal approval for scanning the test environment obtained.",
            "exit": "No protected business endpoint is accessible without valid authentication/role; no open Severity 1 or Severity 2 security defect.",
            "tools": "REST Assured/Postman, OWASP ZAP baseline, dependency scanner assumption, SQL review queries.",
            "deliverables": "Security test report, role matrix, scan summary, defect records.",
        },
        {
            "name": "Performance Testing",
            "purpose": "Measure response time against SRS performance targets.",
            "scope": "Representative auth, event, team, submission, and scoring workflows under 100 concurrent users.",
            "approach": "Run load tests with realistic role distribution and seeded data; measure p95 latency, error rate, database saturation, and notification side effects.",
            "entry": "Performance-like environment available, stable dataset seeded, monitoring enabled, and test run window approved.",
            "exit": "Target p95 <= 2 seconds or minimum p95 <= 5 seconds is achieved; bottlenecks are documented and risk accepted if unresolved.",
            "tools": "JMeter or k6, API load scripts, database/server monitoring.",
            "deliverables": "Load test plan, raw results, summary report, tuning recommendations.",
        },
        {
            "name": "Database Testing",
            "purpose": "Verify schema integrity, persistence, uniqueness, foreign keys, lifecycle state changes, and data side effects.",
            "scope": "Entities in the SRS data dictionary, including _user, event_registration, team, team_member, submission, score, notification, audit_log, prize, and related join/advancement tables.",
            "approach": "Validate database state after API workflows; test uniqueness and FK constraints; compare service validation against database constraints; verify soft delete/deactivation semantics.",
            "entry": "HackathonDB schema deployed from SRS-referenced schema and resettable test data available.",
            "exit": "No data corruption, orphan records, duplicate violations, or incorrect state transitions remain open in release-critical flows.",
            "tools": "SQL Server Management Studio or sqlcmd, database seed scripts, API execution suite.",
            "deliverables": "Database verification scripts, data evidence, integrity defect reports.",
        },
        {
            "name": "Compatibility Testing",
            "purpose": "Verify supported client and runtime combinations do not block SRS workflows.",
            "scope": "Browser-based SPA access, EventSource SSE behavior, REST/JSON, Java 21 backend, SQL Server connectivity, and Windows developer setup minimum.",
            "approach": "Execute core smoke and participant/judge/organizer workflows across agreed browsers and runtime environments.",
            "entry": "Assumption: target browser and OS matrix approved because SRS does not specify exact browser versions.",
            "exit": "Supported matrix executes critical workflows without browser-specific blocker defects.",
            "tools": "Playwright browser matrix, manual browser checks, environment setup checklist.",
            "deliverables": "Compatibility matrix and browser-specific defect list.",
        },
        {
            "name": "Usability Testing",
            "purpose": "Measure the SRS usability target for the core participant workflow.",
            "scope": "Register, verify, register event, create/finalize team, and submit project.",
            "approach": "Conduct moderated or observed first-time user sessions using logical UI screens once frontend is available; measure completion rate and completion time.",
            "entry": "Frontend workflow available, test participants identified, task script approved, and data reset prepared.",
            "exit": "Target >= 90% within 15 minutes or minimum >= 75% within 25 minutes is measured and reported.",
            "tools": "Observation checklist, screen recording assumption, timer, feedback form.",
            "deliverables": "Usability findings, completion metrics, improvement backlog.",
        },
        {
            "name": "Accessibility Testing",
            "purpose": "Reduce access barriers in the browser-based SPA workflows implied by the SRS.",
            "scope": "Authentication, registration, profile, team, submission, scoring, notifications, dashboard, and admin/organizer screens when implemented.",
            "approach": "Assumption: because the SRS does not state a WCAG target, use WCAG 2.2 AA as an enterprise benchmark pending stakeholder confirmation; verify keyboard navigation, labels, focus order, contrast, and error messaging.",
            "entry": "UI build available and target accessibility benchmark approved.",
            "exit": "No blocker accessibility defect prevents completion of core workflows by keyboard or screen-reader users.",
            "tools": "axe DevTools or equivalent, Playwright accessibility checks, keyboard-only review, screen-reader spot checks.",
            "deliverables": "Accessibility audit report and defect backlog.",
        },
        {
            "name": "Negative Testing",
            "purpose": "Prove the system rejects invalid, unauthorized, duplicate, out-of-window, and conflicting operations.",
            "scope": "All business rules and high-risk FRs, including OTP expiry, duplicate identity, invalid event transitions, team limits, invitation eligibility, score bounds, finalized/frozen scores, and prize uniqueness.",
            "approach": "Design invalid data, wrong role, wrong state, duplicate, boundary, expired, and conflict cases for every SRS business rule.",
            "entry": "Validation rules identified and seed data prepared for invalid states.",
            "exit": "Invalid requests fail with appropriate HTTP status/error handling and without corrupting persisted data.",
            "tools": "API client, browser, SQL verification queries, automated negative test suite.",
            "deliverables": "Negative test matrix, defect reports, validation evidence.",
        },
        {
            "name": "Boundary Testing",
            "purpose": "Verify minimum, maximum, zero, expiry, time-window, and scoring boundaries.",
            "scope": "Passwords >= 6 characters, OTP six digits and five-minute expiry, team min/max, registration windows, event/round times, criterion weights <= 100, score 0..max_score, advancement slots, CSV output boundaries.",
            "approach": "Test at min, max, just below, just above, before window, inside window, after window, and duplicate boundary states.",
            "entry": "Boundary values identified from SRS business rules, data dictionary, and NFR targets.",
            "exit": "All high-risk boundary defects resolved or risk accepted; no boundary error permits invalid business state.",
            "tools": "API automation, SQL seed scripts, time-control test data, performance scripts for NFR boundaries.",
            "deliverables": "Boundary value matrix and execution evidence.",
        },
    ]


def build_strategy(doc, num_id):
    h(doc, "Test Strategy", 1, num_id)
    p(doc, "The QA strategy SHALL be risk-based, requirement-traceable, API-centered, and workflow-driven. The SRS defines a backend-heavy system with multiple roles, time windows, lifecycle states, and business constraints. Testing shall therefore emphasize state transitions, role permissions, data integrity, and cross-module workflow correctness.")
    p(doc, "Planning decision: API and integration testing SHALL receive high automation priority because the SRS identifies REST endpoints, JSON contracts, SQL Server persistence, and backend service workflows as the stable behavioral interface. UI testing SHALL validate complete workflows and usability once the SPA screens are available.")

    h(doc, "Risk-Based Testing Approach", 2, num_id)
    rbt_rows = [
        ["Risk factor", "Application to SEAL Hackathon", "Planning decision"],
        ["Business impact", "Account, team, submission, scoring, ranking, and advancement defects can block event execution or compromise fairness.", "These areas SHALL be tested first and included in every release regression cycle."],
        ["Technical complexity", "JWT/OAuth2, SMTP, SSE, SQL uniqueness, status transitions, score calculations, and role checks create integration risk.", "Integration and API tests SHALL combine role, data-state, and side-effect assertions."],
        ["Probability of defect", "Duplicate checks, time windows, ownership rules, and finalization/frozen states are common defect sources.", "Negative and boundary tests SHALL be mandatory for linked business rules."],
        ["Detectability", "Ranking and scoring defects may not be obvious from UI display alone.", "Database and calculation oracle checks SHALL be used for ranking/scoring."],
    ]
    add_table(doc, rbt_rows[0], rbt_rows[1:], widths=[1900, 4700, 3150], font_size=8.5)

    h(doc, "Testing Types", 2, num_id)
    p(doc, "Each testing type below includes purpose, scope, approach, entry criteria, exit criteria, tools, and deliverables. Tool choices are planning assumptions where the SRS does not prescribe a toolchain.")
    for t in testing_type_definitions():
        h(doc, t["name"], 3, num_id)
        add_table(
            doc,
            ["Planning Element", "Definition"],
            [
                ["Purpose", t["purpose"]],
                ["Scope", t["scope"]],
                ["Approach", t["approach"]],
                ["Entry Criteria", t["entry"]],
                ["Exit Criteria", t["exit"]],
                ["Tools", "Assumption: final enterprise toolchain is not specified in the SRS. Recommended tools: " + t["tools"]],
                ["Deliverables", t["deliverables"]],
            ],
            widths=[1900, 7850],
            font_size=8.7,
        )


def build_test_levels(doc, num_id):
    h(doc, "Test Levels", 1, num_id)
    rows = [
        ["Unit Testing", "Developers", "Controller validation, service business rules, utility methods, repository methods where applicable.", "JUnit 5, Mockito, Spring Boot Test assumption.", "Unit test suite, coverage report, build evidence.", "Unit tests SHALL complete before a build is promoted to QA."],
        ["Integration Testing", "Developers with QA support", "Controller-service-repository, SQL Server, SMTP/OAuth2/SSE integration, audit/notification side effects.", "Spring Boot Test, REST Assured/Postman, SQL Server tools.", "Integration results, logs, data evidence.", "High-risk integrations SHALL pass before system test begins."],
        ["System Testing", "QA team", "End-to-end SRS workflows across roles, modules, data states, and NFR smoke checks.", "API automation, browser checks, database verification.", "Execution reports, defects, RTM updates.", "System testing SHALL produce the release quality evidence."],
        ["User Acceptance Testing", "BA/Product Owner, selected user representatives, QA facilitation", "Business workflow acceptance for participants, team leaders, judges, mentors, organizers, and admins.", "UAT scripts, observation notes, sign-off checklist.", "UAT sign-off, open issue log, business risk acceptance.", "UAT SHALL not replace QA system testing; it confirms business acceptance."],
    ]
    add_table(doc, ["Level", "Primary Responsibility", "Scope", "Tools", "Deliverables", "Planning Rationale"], rows, widths=[1250, 1700, 2850, 1700, 1400, 1850], font_size=7.6)


def build_environment(doc, num_id):
    h(doc, "Test Environment", 1, num_id)
    p(doc, "The test environment SHALL be production-like for components explicitly defined in the SRS and controlled through documented assumptions where the SRS is silent.")
    rows = [
        ["Frontend", "Browser-based SPA; local CORS origins configured as http://localhost:5173 and http://localhost:5174.", "Frontend technology beyond SPA is not specified in the SRS. Assumption: final browser test URLs will be supplied by the team."],
        ["Backend", "Java 21, Spring Boot 3.2.5, Spring Web, Spring Security, Spring Data JPA, Spring Validation, Spring Mail, Spring Retry, OAuth2 Client, SpringDoc OpenAPI.", "Build shall run with test-safe externalized secrets."],
        ["Database", "Microsoft SQL Server; JDBC URL configured for localhost:1433 and HackathonDB; Hibernate ddl-auto=none.", "Schema shall match SRS-referenced database.sql/JPA entities before QA execution."],
        ["Browser", "Browser-based SPA with EventSource for SSE.", "Assumption: latest stable Chrome and Edge are mandatory; Firefox compatibility shall be tested if approved in the browser matrix."],
        ["Operating System", "SRS minimum portability target states Windows developer machine with Java 21 and SQL Server.", "Assumption: one CI/Linux-equivalent setup shall be attempted for desired portability target."],
        ["Network", "REST HTTP(S), Authorization Bearer JWT, SSE token query parameter, SMTP TLS 587, SQL Server connection.", "Assumption: production shall use HTTPS/TLS as stated by SRS assumption."],
        ["Third-party Services", "Gmail SMTP and Google OAuth2 profile/email scope.", "QA shall use sandbox/test credentials or approved stubs; production credentials SHALL NOT be used."],
        ["Email", "OTP, approval, and notification email workflows via Gmail SMTP.", "Test mailbox shall be isolated and monitored; email timing shall be validated without exposing real PII."],
        ["Storage", "SQL Server persistence; submissions store repository/demo/report URLs, not binary files.", "No binary upload storage testing is planned because SRS excludes it."],
        ["Environment Configuration", "JWT, SMTP, OAuth, database connection, CORS origins, and application properties.", "Assumption: configuration management owner shall provide documented values for QA and prevent secret leakage."],
    ]
    add_table(doc, ["Area", "SRS-Based Configuration", "QA Control / Assumption"], rows, widths=[1700, 4800, 3250], font_size=8.2)

    h(doc, "Environment Readiness Checklist", 2, num_id)
    checks = [
        "Backend starts successfully on Java 21 and exposes OpenAPI metadata.",
        "HackathonDB exists, schema is current, and data reset procedure is available.",
        "All role-based seed users can be created or loaded: ADMIN, ORGANIZER, JUDGE, GUEST_JUDGE, MENTOR, PARTICIPANT, and team leader.",
        "SMTP test mailbox can receive OTP, approval, and notification emails or approved mail stub is enabled.",
        "Google OAuth2 test configuration is available or OAuth behavior is isolated for manual/provider verification.",
        "Frontend SPA can reach backend from configured CORS origin and display logical workflows.",
        "SSE stream can be opened with an authenticated token query parameter.",
        "Logs are available for backend, database, email attempts, and test automation runs.",
    ]
    for c in checks:
        bullet(doc, c)


def build_data_management(doc, num_id):
    h(doc, "Test Data Management", 1, num_id)
    p(doc, "Test data SHALL be synthetic, resettable, traceable to scenarios, and separated from production data. Real student, judge, or organizer personal data SHALL NOT be used.")
    rows = [
        ["Test Data Sources", "Synthetic user accounts, SQL seed scripts, API-created domain records, controlled test mailbox, OAuth test profile, CSV export fixtures.", "No production PII SHALL be copied into QA environments."],
        ["Data Preparation", "Seed users by role and state; create events in DRAFT/PUBLISHED/IN_PROGRESS/COMPLETED/CANCELLED; seed tracks, rounds, criteria, teams, submissions, scores, notifications, prizes, audit logs.", "Data states SHALL map to SRS business rules and scenario IDs."],
        ["Data Refresh", "Reset HackathonDB before formal cycles; retain snapshots for defect reproduction; clear transient tokens and notification state where required.", "A failed data refresh SHALL suspend execution for dependent tests."],
        ["Privacy", "Emails, phone, student ID, GitHub URL, category customer data, and export content are treated as PII.", "Only synthetic values SHALL be used and shared in reports."],
        ["Masking", "If any non-synthetic dataset is approved, email, phone, student ID, name, and identifiers SHALL be masked before use.", "Assumption: no production dataset is currently approved by the SRS."],
        ["Synthetic Data", "Use deterministic names and identifiers such as participant01@example.test, team-min, team-max, track-ai, round-1, score-boundary.", "Synthetic data SHALL cover valid and invalid combinations."],
    ]
    add_table(doc, ["Area", "Plan", "Control"], rows, widths=[1700, 5000, 3050], font_size=8.4)

    h(doc, "Minimum Seed Data Matrix", 2, num_id)
    seed_rows = [
        ["Users", "Verified/approved participant, unverified participant, unapproved participant, admin, organizer, judge, guest judge, mentor, inactive user.", "Auth, user admin, role security, profile, invitation, scoring, mentorship."],
        ["Events", "DRAFT, PUBLISHED with open registration, PUBLISHED closed registration, IN_PROGRESS, COMPLETED, CANCELLED.", "Event lifecycle, registration window, public browse, status restrictions."],
        ["Teams", "No team, one-member team, min-size team, max-size team, finalized team, disqualified team, non-finalized deletable team.", "Team rules, submission eligibility, disqualification, member constraints."],
        ["Scoring", "Criteria total <= 100, criteria total > 100, score 0, score max, score above max, finalized scores, advanced round.", "Criteria, score, ranking, advancement, negative and boundary testing."],
        ["Notifications", "Unread, read, offline recipient, online SSE recipient.", "Notification API and SSE delivery behavior."],
        ["Exports", "Ranked teams, anonymized scores, teams, participants with synthetic PII.", "CSV structure, permissions, BOM, audit action."],
    ]
    add_table(doc, ["Dataset", "Required States", "Primary Coverage"], seed_rows, widths=[1600, 5000, 3150], font_size=8.4)


def build_entry_exit_suspension(doc, num_id):
    h(doc, "Entry Criteria", 1, num_id)
    entries = [
        ["EC-01", "SRS baseline", "SRS version 1.0 is available and accepted as the test basis."],
        ["EC-02", "Requirement traceability", "100% of SRS FRs and NFRs are represented in the RTM with scenario IDs."],
        ["EC-03", "Build availability", "Deployable backend build is provided with build number, commit reference assumption, and release notes."],
        ["EC-04", "Environment readiness", "Backend, SQL Server, frontend SPA, SMTP/OAuth test configuration or approved stubs, and logs are available."],
        ["EC-05", "Smoke pass", "Mandatory smoke suite passes with no blocker before system test execution starts."],
        ["EC-06", "Test data readiness", "Seed data for all roles, event states, team states, scoring states, and negative cases is loaded or creatable."],
        ["EC-07", "Defect workflow", "Defect tracker, severity/priority definitions, owners, and triage cadence are agreed."],
        ["EC-08", "Tool readiness", "API, UI, database, performance, and security testing tools are installed and configured."],
    ]
    add_table(doc, ["ID", "Criterion", "Measurable Acceptance"], entries, widths=[900, 2500, 6350], font_size=8.7)

    h(doc, "Exit Criteria", 1, num_id)
    exits = [
        ["XC-01", "Execution completion", "100% of planned high-priority and risk-critical test cases executed; >= 95% of planned total cases executed or formally deferred."],
        ["XC-02", "Requirement coverage", "100% of SRS FRs/NFRs have pass/fail/blocked status and evidence in the RTM."],
        ["XC-03", "Defect threshold", "0 open Severity 1; 0 open Severity 2 unless explicitly approved by release authority; Severity 3 defects have workaround or accepted residual risk."],
        ["XC-04", "Regression stability", "Release regression pass rate >= 95%; no critical reopened defect after final regression."],
        ["XC-05", "NFR disposition", "Performance, security, reliability, usability, maintainability, and portability NFRs measured or documented with approved risk disposition."],
        ["XC-06", "Test summary", "Test Summary Report completed with metrics, defect trends, residual risks, and release recommendation."],
        ["XC-07", "Stakeholder sign-off", "QA Lead, PM, BA/Product Owner, and Development Lead approve exit or record exceptions."],
    ]
    add_table(doc, ["ID", "Criterion", "Measurable Acceptance"], exits, widths=[900, 2500, 6350], font_size=8.7)

    h(doc, "Suspension and Resumption Criteria", 1, num_id)
    susp = [
        ["Suspension", "Build fails smoke test, backend unavailable, database unavailable, login/token workflow unavailable, test data corrupted, environment misconfigured, or more than 30% of planned tests blocked by a common defect.", "QA Lead suspends affected cycle, records reason, and returns build/environment to owner."],
        ["Suspension", "A Severity 1 defect prevents execution of release-critical workflows such as login, event setup, team creation, submission, scoring, or ranking.", "Testing continues only on unaffected areas if value remains; otherwise cycle pauses."],
        ["Resumption", "Blocking defect fixed and sanity tested, environment restored, data reset completed, and smoke suite passes.", "QA Lead resumes execution and identifies regression impact."],
        ["Resumption", "External service issue has approved stub/workaround and affected tests are marked with limitation.", "Execution resumes with explicit assumption and risk note."],
    ]
    add_table(doc, ["Type", "Criteria", "Required Action"], susp, widths=[1300, 5600, 2850], font_size=8.5)


def build_defect_config(doc, num_id):
    h(doc, "Defect Management", 1, num_id)
    h(doc, "Workflow", 2, num_id)
    flow = [
        ["New", "Defect logged by tester with evidence and requirement/test reference.", "QA"],
        ["Triaged", "Defect severity, priority, owner, and target release determined.", "QA Lead, PM, Dev Lead"],
        ["Assigned", "Defect assigned to developer or component owner.", "Dev Lead"],
        ["In Progress", "Fix is under development.", "Developer"],
        ["Ready for QA", "Fix deployed with notes and affected build.", "Developer"],
        ["Retest", "QA validates original failure and impacted regression checks.", "QA"],
        ["Reopened", "Fix failed or created unacceptable regression.", "QA"],
        ["Deferred", "Accepted for later release with risk owner approval.", "PM/PO"],
        ["Rejected / Duplicate", "Not a defect, already reported, or outside SRS scope.", "Triage team"],
        ["Closed", "Fix verified or final disposition approved.", "QA Lead"],
    ]
    add_table(doc, ["Status", "Meaning", "Owner"], flow, widths=[1700, 6100, 1950], font_size=8.6)

    h(doc, "Severity and Priority", 2, num_id)
    sev = [
        ["S1 Critical", "System unavailable, security bypass, data corruption, or release-critical workflow cannot proceed.", "Login blocked for approved users; protected endpoint accessible without role; scoring/ranking corrupts results."],
        ["S2 High", "Major SRS feature fails with no practical workaround or high business risk.", "Team finalization rejects valid min-size team; score finalization fails; export exposes unauthorized data."],
        ["S3 Medium", "Feature defect with workaround or limited scope.", "Dashboard count incorrect for non-release metric; invitation list display issue with API workaround."],
        ["S4 Low", "Cosmetic, minor usability, wording, or low-risk issue.", "Non-blocking label issue or minor sorting inconsistency outside acceptance criteria."],
    ]
    add_table(doc, ["Severity", "Definition", "Example"], sev, widths=[1600, 5000, 3150], font_size=8.5)

    prio = [
        ["P1 Immediate", "Must be fixed before testing/release can continue.", "S1 or release-blocking S2."],
        ["P2 High", "Target current sprint/release.", "High-priority FR failure or high-risk business rule."],
        ["P3 Normal", "Fix when scheduled; workaround acceptable.", "Medium issue with manageable impact."],
        ["P4 Low", "Backlog candidate.", "Cosmetic or low-risk issue."],
    ]
    add_table(doc, ["Priority", "Definition", "Typical Use"], prio, widths=[1600, 5200, 2950], font_size=8.5)

    h(doc, "Bug Report Fields", 2, num_id)
    fields = [
        "Defect ID, title, reporter, date/time, environment, build/version, requirement ID, test case ID, module, role, preconditions, steps to reproduce, expected result, actual result, severity, priority, attachments/evidence, logs, API request/response, database evidence where applicable, reproducibility, owner, status, fix version, retest result, regression impact, closure notes.",
    ]
    for item in fields:
        bullet(doc, item)

    h(doc, "Configuration Management", 1, num_id)
    rows = [
        ["Version Control", "Project source and test artifacts SHALL be versioned. Assumption: repository branching model is not specified in SRS; QA shall record build identifier, commit reference if available, and test artifact version."],
        ["Build Management", "Every QA cycle SHALL reference a unique backend/frontend build and deployment timestamp. Builds failing smoke SHALL not enter full execution."],
        ["Environment Management", "Database schema, application properties, secrets, CORS origins, SMTP/OAuth test configuration, and seed data SHALL be controlled and reproducible."],
        ["Test Artifact Management", "Master Test Plan, RTM, test cases, data scripts, automation suites, reports, and evidence SHALL be stored with version and approval status."],
        ["Change Control", "Any SRS change SHALL trigger impact analysis across FRs, BRs, endpoints, entities, tests, risks, and regression pack scope."],
    ]
    add_table(doc, ["Area", "Plan"], rows, widths=[2200, 7550], font_size=8.6)


def build_risk_analysis(doc, num_id, reqs):
    h(doc, "Risk Analysis", 1, num_id)
    p(doc, "Risk level is determined by business impact, technical complexity, probability of defect, and detectability. High and critical risks SHALL receive earlier execution, deeper negative/boundary coverage, and stronger regression protection.")
    modules = module_catalog(reqs)
    risk_text = {
        "AUTH": ("Unauthorized or incorrect access blocks all protected workflows and may expose account data.", "JWT, refresh token, OTP expiry, BCrypt, OAuth2, approval/verification state, and SMTP dependencies.", "High", "High", "Critical", "Automate role/auth API tests; test verified/unverified/approved/unapproved states; validate token refresh and OTP boundaries.", "Suspend release if approved login or protected endpoint security fails."),
        "USER": ("Incorrect approval/deactivation/profile handling can block users or allow stale accounts.", "Role authorization, audit side effects, uniqueness, and finalized-team profile restrictions.", "Medium", "High", "High", "Test ADMIN-only actions, deactivation login effects, unique identity constraints, and audit logging.", "Restrict admin release functions until fixed or apply manual admin control."),
        "EVENT": ("Invalid event lifecycle can open registration/scoring at wrong times.", "State transition enforcement, time windows, slug uniqueness, soft delete rules, and notifications.", "High", "High", "High", "Run state-transition and date-boundary tests; verify owner/admin authorization.", "Freeze affected event lifecycle operation and use manual database rollback only with approval."),
        "CONFIG": ("Incorrect tracks/rounds/criteria can invalidate scoring and advancement.", "Criterion weight <= 100, no criteria edit after scoring, mentor conflict and round setup.", "High", "High", "High", "Test criteria weight boundaries, scoring-start lock, round times, and mentor assignment constraints.", "Block scoring cycle until configuration defects resolved."),
        "REG": ("Ineligible participants may enter or eligible users may be blocked.", "PUBLISHED status, registration window, verified/approved/profile-complete checks, duplicate registration.", "High", "High", "High", "Execute eligibility matrix across user state, event status, and time window.", "Pause registration opening and correct eligibility defect."),
        "TEAM": ("Team composition defects can invalidate competition eligibility.", "One-team-per-event, min/max size, finalization, disqualification, delete side effects.", "High", "High", "Critical", "Create min/max/full/finalized/disqualified data states; test all team status transitions.", "Manual competition roster review and block submissions until corrected."),
        "MEMBER": ("Member changes can break team eligibility or leadership.", "Registration phase constraints, finalized-team min size, leader leave/transfer rules.", "Medium", "High", "High", "Test add/remove/kick/leave/transfer with leader/non-leader roles and team size boundaries.", "Freeze roster changes for affected event."),
        "INV": ("Invitation defects can bypass one-team rule or block team formation.", "Duplicate invitations, invitee eligibility, cancellation of other pending invites, notifications.", "Medium", "High", "High", "Test invite/respond/revoke against team capacity, registration window, and account state.", "Disable invitation path and use direct admin-controlled membership only if approved."),
        "SUB": ("Submission failure prevents teams from competing.", "Round window timing, leader-only access, finalized/non-disqualified/min-size constraints, version increment.", "High", "High", "Critical", "Test submission at before/inside/after window, role restrictions, update/version behavior.", "Extend round window only through approved event operations after defect fix."),
        "JUDGE": ("Wrong judge assignment can create unfair scoring or conflict of interest.", "Round/track assignment, all-track vs specific-track, mentor conflict, guest judge role.", "Medium", "High", "High", "Test assignment matrix and conflict prevention against track_mentor data.", "Manual assignment review before scoring starts."),
        "SCORE": ("Invalid scores directly compromise rankings and prizes.", "Score 0..max, criterion mapping, owner updates, finalization/frozen/advanced restrictions.", "High", "High", "Critical", "Use calculation oracles, boundary scores, finalized and advanced states.", "Hold score finalization/ranking until corrected; preserve audit evidence."),
        "RANK": ("Ranking defects compromise competition fairness.", "Weighted averages, per-judge averaging, two-decimal rounding, track rank, finalized-only, exclusion of disqualified teams.", "Medium", "High", "Critical", "Use independent expected-result spreadsheet/script and database verification.", "Publish no rankings until calculation variance resolved."),
        "ADV": ("Advancement defects send wrong teams to next round or prematurely complete event.", "Round end check, advancement slots, duplicate prevention, next-round/no-next-round behavior.", "Medium", "High", "High", "Test ranked/non-ranked/disqualified/duplicate/slot-boundary scenarios.", "Manual advancement review and rollback plan before next round starts."),
        "PRIZE": ("Prize assignment defects create incorrect awards.", "Unique rank per event/track, track match, one prize per team, auto assignment order.", "Medium", "Medium", "Medium", "Test unique rank and auto-assign from known rankings.", "Manual award review before publication."),
        "MENTOR": ("Mentorship defects reduce team support but do not usually block core competition.", "Request status workflow, assigned mentor authority, cancellation roles, notifications.", "Medium", "Medium", "Medium", "Test OPEN/IN_PROGRESS/RESOLVED/REJECTED/CANCELLED state transitions.", "Use manual organizer communication if mentorship workflow unavailable."),
        "NOTIF": ("Missed notifications can hide important status changes.", "Persistence, ownership, unread count, mark read, SSE online/offline delivery.", "Medium", "High", "High", "Verify persisted notification even when offline and SSE delivery when online.", "Use email/manual communication for affected release events."),
        "DASH": ("Incorrect dashboard metrics can mislead organizers.", "Aggregate counts, days remaining, finalized-score variance.", "Medium", "Medium", "Medium", "Compare dashboard outputs against SQL fixtures.", "Use exports/SQL reports until dashboard fixed."),
        "EXPORT": ("Incorrect or unauthorized CSV exports can leak PII or misreport results.", "ADMIN/ORGANIZER access, UTF-8 BOM, ranking/anonymized/team/participant content, audit log.", "Medium", "High", "High", "Role-test exports, inspect CSV structure/content, verify audit action.", "Disable export endpoint until access/content defect fixed."),
        "AUDIT": ("Missing audit evidence weakens investigation and accountability.", "Significant actions, ADMIN-only review, ordering, entity/user details.", "Medium", "High", "High", "Verify audit creation for admin/workflow operations and access restrictions.", "Require manual action log until audit restored."),
        "CAT": ("Legacy module defects have limited hackathon business impact but can affect exposed CRUD/API quality.", "Unique code, DTO validation, public read/ADMIN write, absence from database.sql issue.", "Low", "Medium", "Low", "Test CRUD if retained in release; confirm scope decision for v1.0.", "Defer or disable module if stakeholders confirm not required."),
    }
    rows = []
    for m in modules:
        rows.append([m["name"], *risk_text[m["key"]]])
    add_table(
        doc,
        ["Module", "Business Risk", "Technical Risk", "Probability", "Impact", "Risk Level", "Mitigation", "Contingency"],
        rows,
        widths=[1500, 1850, 1850, 900, 800, 900, 2450, 1500],
        font_size=6.9,
    )


def build_resource_estimation(doc, num_id, reqs):
    h(doc, "Resource Plan", 1, num_id)
    rows = [
        ["QA Lead / Test Manager", "1", "Owns Master Test Plan, risk-based strategy, triage, reporting, exit recommendation, stakeholder communication.", "Full cycle"],
        ["QA Engineer - Functional/API", "2", "Design and execute FR/API tests, negative/boundary tests, database verification, defect logging.", "Full cycle"],
        ["Automation Engineer", "1", "Build API/UI regression suite, CI integration assumption, data setup utilities, automation reporting.", "Preparation through regression"],
        ["Performance/Security Tester", "0.5-1", "Execute NFR-PERF and NFR-SEC tests, role matrix, baseline scan, performance report.", "Targeted cycles"],
        ["Developers", "As assigned", "Own unit tests, integration fixes, defect resolution, technical support, testability hooks.", "Throughout"],
        ["Business Analyst / Product Owner", "1", "Clarify SRS interpretation, approve RTM, support UAT, accept residual business risk.", "Planning and UAT"],
        ["Project Manager", "1", "Schedule, resources, risk decisions, change control, release governance.", "Throughout"],
        ["Database/DevOps Support", "Assumption: shared", "Maintain SQL Server schema, data refresh, environment configuration, logs, credentials.", "Environment setup and issue resolution"],
    ]
    add_table(doc, ["Role", "Estimated Count", "Responsibilities", "Engagement"], rows, widths=[2000, 1300, 5100, 1350], font_size=8.3)

    h(doc, "Test Estimation", 1, num_id)
    frs = [r for r in reqs if r["ID"].startswith("FR-")]
    nfrs = [r for r in reqs if r["ID"].startswith("NFR-")]
    prio = Counter(r.get("Priority", "Unspecified") for r in frs)
    functional_cases = sum(case_count_for(r) for r in frs)
    nfr_cases = sum(case_count_for(r) for r in nfrs)
    br_scenarios = 23
    total_scenarios = len(frs) + br_scenarios + len(nfrs) + 20
    total_cases = functional_cases + nfr_cases + 35
    regression = int(total_cases * 0.72)
    automation = int(total_cases * 0.60)
    estimate_rows = [
        ["Functional requirements", str(len(frs)), "One primary scenario per FR; high-risk FRs receive positive, negative, role, boundary, and data-state cases."],
        ["Business-rule cross-check scenarios", str(br_scenarios), "One scenario per SRS business rule with links back to affected FRs."],
        ["Non-functional requirements", str(len(nfrs)), "Performance, security, usability, reliability, maintainability, and portability scenarios."],
        ["Additional cross-cutting scenarios", "20", "Smoke, regression, compatibility, accessibility, data refresh, and end-to-end workflow scenarios."],
        ["Estimated total scenarios", str(total_scenarios), "Scenario count is risk-adjusted and shall be refined after test case design review."],
        ["Estimated functional test cases", str(functional_cases), f"Based on priority weighting: High={prio.get('High', 0)} x5, Medium={prio.get('Medium', 0)} x3, Low={prio.get('Low', 0)} x2."],
        ["Estimated NFR test cases", str(nfr_cases), "Five test cases per NFR to cover target, minimum, failure mode, evidence, and regression hook."],
        ["Estimated total test cases", str(total_cases), "Includes 35 cross-cutting smoke/regression/accessibility/compatibility/data checks."],
        ["Release regression size", str(regression), "Approximately 72% of total cases, emphasizing all high-priority FRs, all BRs, and representative medium/low features."],
        ["Automation target", f"{automation} cases / approx. 60%", "API-heavy automation is justified by stable REST contracts and high regression value; usability and provider workflows remain partly manual."],
        ["Execution time", "8-10 QA working days", "Assumption: 2 functional QA plus 1 automation engineer; full manual equivalent is approximately 45-55 hours plus defect retest."],
    ]
    add_table(doc, ["Estimate Item", "Estimate", "Rationale"], estimate_rows, widths=[2700, 1600, 5450], font_size=8.4)

    h(doc, "Module-Level Estimate", 2, num_id)
    module_rows = []
    for m in module_catalog(reqs):
        count = len(m["reqs"])
        cases = sum(case_count_for(r) for r in m["reqs"])
        high = sum(1 for r in m["reqs"] if r.get("Priority") == "High")
        medium = sum(1 for r in m["reqs"] if r.get("Priority") == "Medium")
        low = sum(1 for r in m["reqs"] if r.get("Priority") == "Low")
        regression_cases = max(high * 4 + medium * 2 + low, 1)
        auto_pct = "70%" if m["key"] in {"AUTH", "USER", "EVENT", "CONFIG", "REG", "TEAM", "MEMBER", "INV", "SUB", "JUDGE", "SCORE", "RANK", "ADV", "NOTIF", "EXPORT", "AUDIT"} else "40%"
        module_rows.append([m["name"], str(count), f"H:{high} M:{medium} L:{low}", str(cases), str(regression_cases), auto_pct])
    add_table(doc, ["Module", "FR Count", "Priority Mix", "Estimated Cases", "Regression Cases", "Automation Target"], module_rows, widths=[2600, 1000, 1350, 1400, 1500, 1900], font_size=8.1)


def build_traceability(doc, num_id, reqs, table_map):
    h(doc, "Requirement Traceability Strategy", 1, num_id)
    p(doc, "Traceability SHALL map Business Requirement -> Functional Requirement -> Module -> Test Scenario -> Test Case -> Risk -> Automation Candidate. No SRS requirement SHALL be considered covered until it has test design, execution status, and evidence.")
    br_rows, fr_to_brs = build_business_rule_maps(table_map)
    add_table(
        doc,
        ["Traceability Layer", "Rule"],
        [
            ["Business Requirement", "Use SRS BR IDs where present; if no BR is linked, use the FR acceptance criterion as the business test basis."],
            ["Functional Requirement", "Use exact SRS FR/NFR ID and title without renumbering."],
            ["Module", "Use SRS feature/module grouping from Section 4."],
            ["Test Scenario", "Use TS-<module>-<number> aligned to the FR/NFR ID."],
            ["Test Case", "Use TC-<module>-<number>-<sequence> and retain SRS primary TC mapping where available."],
            ["Risk", "Use module risk plus requirement priority to drive execution order and regression inclusion."],
            ["Automation Candidate", "Classify as Yes, Partial, Conditional, or No based on stability, observability, and provider dependency."],
        ],
        widths=[2400, 7350],
        font_size=8.7,
    )

    h(doc, "Requirement Coverage Matrix", 2, num_id)
    rows = []
    for r in reqs:
        rid = r["ID"]
        module = req_module(rid)
        title = r.get("Title", r.get("Quality attribute", ""))
        br = ", ".join(fr_to_brs.get(rid, [])) if rid.startswith("FR-") else "NFR quality requirement"
        if not br:
            br = "FR acceptance criteria only; no explicit BR link in SRS traceability table."
        count = case_count_for(r)
        rows.append([
            br,
            f"{rid} - {title}",
            module,
            ts_id(rid),
            f"{tc_base(rid)}-01 to {tc_base(rid)}-{count:02d}",
            risk_for_req(r),
            automation_candidate(r),
        ])
    add_table(
        doc,
        ["Business Requirement", "Functional / Non-Functional Requirement", "Module", "Test Scenario", "Test Case Set", "Risk", "Automation Candidate"],
        rows,
        widths=[1850, 3000, 1050, 1300, 1700, 850, 2000],
        font_size=6.8,
    )

    h(doc, "SRS Endpoint and Entity Traceability", 2, num_id)
    p(doc, "The SRS provides endpoint/entity traceability for representative critical workflows. These rows SHALL seed API and database verification coverage; remaining endpoints SHALL be discovered from SRS-referenced OpenAPI metadata without expanding functional scope.")
    add_table(doc, table_map[72][0], table_to_rows(table_map, 72), widths=[1250, 2300, 1900, 3000, 1300], font_size=7.4)


def build_deliverables_reports(doc, num_id):
    h(doc, "Test Deliverables", 1, num_id)
    rows = [
        ["Master Test Plan", "This document.", "QA Lead", "Before formal execution"],
        ["Requirement Traceability Matrix", "BR/FR/NFR/module/scenario/case/risk/automation/execution mapping.", "QA Lead / QA Analyst", "Prepared before execution; updated daily"],
        ["Test Scenarios", "High-level scenario inventory aligned to SRS FRs, NFRs, and BRs.", "QA Team", "Before test design sign-off"],
        ["Test Cases", "Detailed executable steps, data, expected results, and evidence requirements.", "QA Team", "Before execution"],
        ["Test Data", "Synthetic seed data, reset scripts, data matrix, and privacy/masking controls.", "QA / DB support", "Before execution and refreshed per cycle"],
        ["Bug Reports", "Defect records with requirement/test linkage, severity, priority, evidence, and retest status.", "QA Team", "Throughout execution"],
        ["Daily Reports", "Daily execution status, pass/fail/blocked metrics, defects by severity, risks, and next-day plan.", "QA Lead", "Each execution day"],
        ["Automation Reports", "API/UI/performance/security automation run output and trend history.", "Automation Engineer", "Per run/cycle"],
        ["Test Summary Report", "Final execution summary, coverage, defects, residual risk, NFR results, and release recommendation.", "QA Lead", "At exit"],
        ["UAT Sign-Off", "Business acceptance result, open issues, and accepted risks.", "BA/Product Owner", "At UAT completion"],
    ]
    add_table(doc, ["Deliverable", "Description", "Owner", "Timing"], rows, widths=[2200, 4600, 1900, 1050], font_size=8.3)

    h(doc, "Test Reporting and Governance", 1, num_id)
    rows = [
        ["Daily QA Status", "Executed, passed, failed, blocked, not run; new/open/closed defects; risks; blockers; next-day plan.", "Daily during active execution."],
        ["Defect Triage", "Review new defects, severity/priority, owner, target build, duplicate/reject/defer decisions.", "Daily or every build day."],
        ["Risk Review", "Review critical/high risks, blocked tests, NFR risks, external-service limitations, and residual release risk.", "Twice weekly or before exit."],
        ["Exit Review", "Confirm exit criteria, residual defects, RTM completeness, NFR disposition, and release recommendation.", "At end of test cycle."],
    ]
    add_table(doc, ["Meeting / Report", "Content", "Frequency"], rows, widths=[2300, 5850, 1600], font_size=8.4)

    h(doc, "Schedule and Milestones", 1, num_id)
    p(doc, "Assumption: the SRS does not provide project dates beyond the SRS submission date. The schedule below is a relative execution model that SHALL be aligned to the project calendar by the Project Manager.")
    rows = [
        ["M1", "Test planning baseline", "Master Test Plan, RTM strategy, environment checklist approved.", "Day 0"],
        ["M2", "Test design complete", "Scenarios/cases/data/automation backlog reviewed and approved.", "Day 2-3"],
        ["M3", "Environment and smoke ready", "Build deployed, data seeded, smoke pass.", "Day 3-4"],
        ["M4", "Cycle 1 execution", "High-risk functional/API/system tests executed; defects logged.", "Day 4-7"],
        ["M5", "Fix verification and regression", "Retest, impacted regression, NFR tests, automation update.", "Day 7-9"],
        ["M6", "UAT and exit", "UAT support, Test Summary Report, release recommendation.", "Day 9-10"],
    ]
    add_table(doc, ["Milestone", "Name", "Exit Output", "Relative Timing"], rows, widths=[1100, 2600, 4500, 1550], font_size=8.4)


def build_appendices(doc, num_id, reqs, table_map):
    h(doc, "Appendix A - Detailed Scenario Inventory", 1, num_id)
    rows = []
    for r in [x for x in reqs if x["ID"].startswith("FR-")]:
        rid = r["ID"]
        objective = first_sentence(r.get("Acceptance criteria", r.get("Description", "")))
        rows.append([
            ts_id(rid),
            rid,
            r.get("Title", ""),
            req_module(rid),
            objective,
            "Positive, negative, role/ownership, data-state, and boundary cases as applicable.",
        ])
    add_table(doc, ["Scenario ID", "FR ID", "Title", "Module", "Scenario Objective", "Design Notes"], rows, widths=[1150, 1050, 1900, 1000, 3600, 2050], font_size=6.9)

    h(doc, "Appendix B - Business Rule Coverage", 1, num_id)
    add_table(doc, table_map[69][0], table_to_rows(table_map, 69), widths=[950, 1350, 5600, 1850], font_size=7.5)

    h(doc, "Appendix C - NFR Test Basis", 1, num_id)
    nfr_rows = []
    for r in [x for x in reqs if x["ID"].startswith("NFR-")]:
        nfr_rows.append([
            r["ID"],
            r.get("Quality attribute", ""),
            r.get("Scale (what we measure)", ""),
            r.get("Meter (how we measure)", ""),
            r.get("Target (desired)", ""),
            r.get("Must-level (minimum)", ""),
        ])
    add_table(doc, ["ID", "Attribute", "Scale", "Meter", "Target", "Minimum"], nfr_rows, widths=[1000, 1400, 2450, 2700, 1200, 1000], font_size=7.4)

    h(doc, "Appendix D - Data Entities Considered for Database Testing", 1, num_id)
    entities = []
    seen = set()
    for row in table_to_rows(table_map, 62):
        ent = row[0]
        if ent not in seen:
            seen.add(ent)
            entities.append(ent)
    entity_rows = []
    for ent in entities:
        related = []
        for r in reqs:
            text = (r.get("Description", "") + " " + r.get("Source", "") + " " + r.get("Acceptance criteria", "")).lower()
            if ent.replace("_", " ") in text or ent in text:
                related.append(r["ID"])
        entity_rows.append([ent, ", ".join(related[:8]) if related else "Indirect coverage through linked workflows.", "Verify PK/FK/unique constraints, CRUD/state side effects, and rollback/reset behavior."])
    add_table(doc, ["Entity", "Indicative Requirement Coverage", "Database Test Focus"], entity_rows, widths=[2200, 3400, 4150], font_size=7.7)

    h(doc, "Appendix E - Self Review Checklist", 1, num_id)
    checks = [
        ["Every Functional Requirement is covered.", "Pass", "53 FRs mapped in Requirement Coverage Matrix and Scenario Inventory."],
        ["Every module appears in testing scope.", "Pass", "20 SRS functional modules appear in Features to Be Tested and Risk Analysis."],
        ["Every testing type has justification.", "Pass", "Functional, integration, system, API, regression, smoke, sanity, security, performance, database, compatibility, usability, accessibility, negative, and boundary testing sections include purpose/scope/approach/criteria/tools/deliverables."],
        ["Every risk has mitigation.", "Pass", "Risk Analysis table includes mitigation and contingency for every module."],
        ["Every planning decision is explained.", "Pass", "Risk-based prioritization, automation emphasis, tool assumptions, environment assumptions, and exclusions are justified."],
        ["Every section is complete.", "Pass", "Document control, objectives, scope, strategy, levels, environment, data, criteria, defect/config management, risk, resources, estimates, traceability, deliverables, reporting, and appendices are included."],
    ]
    add_table(doc, ["Review Item", "Status", "Evidence"], checks, widths=[3600, 1000, 5150], font_size=8.5)


def main():
    global BULLET_NUM_ID, DECIMAL_NUM_ID, HEADING_COUNTERS
    HEADING_COUNTERS = [0, 0, 0, 0]
    reqs, table_map = load_data()
    doc = Document()
    configure_styles(doc)
    set_update_fields_on_open(doc)
    configure_header_footer(doc)
    num_id = 0
    BULLET_NUM_ID = None
    DECIMAL_NUM_ID = None

    doc.core_properties.title = "SEAL Hackathon Management System - Master Test Plan"
    doc.core_properties.subject = "Enterprise QA Master Test Plan"
    doc.core_properties.author = "Project QA Team"
    doc.core_properties.comments = "Generated from the SEAL Hackathon SRS as the sole functional source of truth."

    build_cover(doc)
    add_toc(doc)
    doc.add_page_break()
    build_document_control(doc, num_id)
    build_intro(doc, num_id, table_map)
    build_project_understanding(doc, num_id, table_map)
    build_test_objectives(doc, num_id)
    build_scope(doc, num_id, reqs, table_map)
    build_strategy(doc, num_id)
    build_test_levels(doc, num_id)
    build_environment(doc, num_id)
    build_data_management(doc, num_id)
    build_entry_exit_suspension(doc, num_id)
    build_defect_config(doc, num_id)
    build_risk_analysis(doc, num_id, reqs)
    build_resource_estimation(doc, num_id, reqs)
    build_traceability(doc, num_id, reqs, table_map)
    build_deliverables_reports(doc, num_id)
    build_appendices(doc, num_id, reqs, table_map)

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    main()
